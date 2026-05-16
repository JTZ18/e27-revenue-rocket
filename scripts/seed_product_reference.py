"""Parse product reference docx into KB markdown via LLM."""
import asyncio
from datetime import date
from pathlib import Path

import docx

from intel_engine.llm.client import LLMClient, LLMProvider
from intel_engine.schemas.kb import KBDomain, KBEntry, KBFrontmatter
from scripts.seed_faq import slugify

PARSE_PROMPT = (
    "You are converting a Boldr product reference document into a structured list "
    "of product entries.\n\n"
    "Each entry is one product (watch model, strap, accessory) OR a related "
    "catalogue table (e.g., strap compatibility).\n\n"
    'Return JSON: {"entries": [{"sku": "BOLDR-XXX", "name": "...", "specs": "..."}, ...]}\n\n'
    "Rules:\n"
    '- "sku" should be the most distinctive identifier in the source. If absent, '
    "use a kebab-case slug of the product name.\n"
    '- "name" is the human-readable product name.\n'
    '- "specs" is a verbatim chunk of all relevant detail: materials, dimensions, '
    "water resistance, movement, compatibility notes, warnings.\n"
    "- Preserve units, standards (ISO 3157, Grade 5 Ti, EU REACH), and any ⚠ callouts.\n"
    '- Treat the strap catalogue / Q-A quick-reference table as its own entry with '
    'sku "STRAP-CATALOGUE".'
)


def extract_docx_text(docx_path: Path) -> str:
    doc = docx.Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


async def parse_with_llm(text: str) -> list[dict[str, str]]:
    client = LLMClient(provider=LLMProvider.kimi)
    result = await client.complete_json(
        system=PARSE_PROMPT,
        user=f"Document:\n\n{text}",
    )
    return result["entries"]


def write_products_from_parsed(
    parsed: list[dict[str, str]],
    out_dir: Path,
    source_file: str,
) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for item in parsed:
        sku = item["sku"]
        slug = slugify(sku) if sku else slugify(item["name"])
        entry = KBEntry(
            frontmatter=KBFrontmatter(
                slug=slug,
                title=item["name"],
                domain=KBDomain.spec,
                themes=["product_general"],
                sources=[f"data/{source_file}"],
                last_verified=date.today(),
            ),
            body=item["specs"],
        )
        path = out_dir / f"{slug}.md"
        path.write_text(entry.to_markdown())
        written.append(path)
    return written


async def main() -> None:
    src = Path("data/05b_product_reference.docx")
    text = extract_docx_text(src)
    parsed = await parse_with_llm(text)
    written = write_products_from_parsed(parsed, Path("kb/products"), source_file=src.name)
    print(f"Wrote {len(written)} product entries to kb/products/")


if __name__ == "__main__":
    asyncio.run(main())
